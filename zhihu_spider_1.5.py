#作者：CapGhost
'''更新说明：v1.2 把图片的属性由data-original改为data-actualsrc
			 v1.3 文本通过RichText属性来获取，匹配不再用BS来处理标签，全部用正则处理
			 v1.4 后缀名采用图源的后缀，添加转换为word功能
			 v1.5 识别收藏内有效链接，按内容标题生成文件夹，修复了同一个问题下两个答案保存混乱的问题

	TODO：
		  封装
	潜在bug：发现两个。
	1.部分知乎文章的图源标签混乱，会导致图片数量不对。
	2.部分回答（似乎是敏感回答）会导致登录后才能查看。由于本代码没有读取登录cookie，而且知乎在2017年10月左右，
	更改了UA验证，导致无法提交正确的post包。需要破解大神的支援（对比之下验证码的破解似乎不难）

	至于标点符号啊、<code>标签、生成word文档的正文字体啊、知乎文章题图不进行爬取啊等等细节。。。实在不想处理了，太麻烦了 

	新bug：
'''

import requests
import time
from bs4 import BeautifulSoup
import os
import re
from docx import Document
from docx.shared import Inches


def html_download(url):
	agent='Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/51.0.2704.106 Safari/537.36'
	headers={'User-Agent':agent}

	try:
		resp=requests.get(url,headers=headers)
		resp.raise_for_status()
		resp.encoding=resp.apparent_encoding
		return resp.text
	except:
		return "网页获取失败"	

def get_pic_url(content):		#解析网页中图片的url，content形参为html.text
	soup=BeautifulSoup(content,"html.parser")
	data=soup.find_all('img')
	url_list=[]	
	for i in data:
		try:
			if 'data-original' in i.attrs:								#知乎里图片有时候只有data-original属性，有时候只有data-actualsrc属性，有时候两个都有。。。我也很头疼怎么处理才完美
				url_list.append(i.attrs['data-original'])				#只有data-actualsrc属性的图片，貌似都不重要
#			if 'data-actualsrc' in i.attrs:
#				url_list.append(i.attrs['data-actualsrc'])	
		except:
			continue	#<img>里非法的属性全忽略
			
	return url_list


def save_pic(url_list,root):		#通过url把图片下载到本地
	agent='Mozilla/5.0 (Windows NT 6.1; WOW64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/51.0.2704.106 Safari/537.36'
	headers={'User-Agent':agent}		

	try:
		if not os.path.exists(root):
			os.makedirs(root)									#多级目录是makedirs，单级目录是mkdir

		for i in range(len(url_list)//2):
			pic=requests.get(url_list[i*2],headers=headers) 	#url_list[偶数]为图片源地址，知乎的H5会将图片显示两次，只留一次的
			suffix=url_list[i*2].split('.')[-1]					#获取图片后缀名，方便处理gif等不同类型图片
			path=root+str(i)+'.'+suffix

			with open(path,'wb') as f:
				f.write(pic.content)
				f.close()
	except:
		print ('图片保存失败')
	return print('图片保存完毕')



def get_text(content):		#解析网页中的文本，content形参为html.text
	soup=BeautifulSoup(content,"html.parser")
	data=soup.div.find_all(attrs='RichText')
	temp=[]		#原始正文数据
	text=[]		#处理后的正文数据

	cnt=0		#图片计数

#	pat_text=r'\>[^\<]+\<|data-original=".*?"'							#	>文字内容< 或者 data-actualsrc="https://pic1.zhimg.com/v2-d69c8bcbc07db60932a5ff0be4fb16ea_b.jpg"
	pat_text=r'\>[^\<]+\<|<img(.*?)>'
	pat_picurl=r'data-original=".*?"'
	temp=re.findall(pat_text,str(data))

	for t in temp:
		if 'zh-lightbox-thumb lazy' in t and re.search(pat_picurl,t):
			text.append('图片'+str(cnt)+'\n')
			cnt=cnt+1
		else:
			t_re=t.replace('>','').replace('<','\n')	#只是取消">"<"这两个标识符，链式replace即可
			text.append(t_re)
	text_str = "".join(text)		
	return text_str			

def get_title(content):		#解析网页中的标题，content形参为html.text
	soup=BeautifulSoup(content,"html.parser")
	answer_num=soup.find('a',attrs='href')
	return soup.find('title').text


def save_text(text,root):

#	try:
	if not os.path.exists(root):
		os.makedirs(root)


	path=root+'context.txt'		

	with open(path,'w',encoding='utf-8') as f:
		f.write(str(text))
		f.close()
#	except:
#		print ('文本保存失败')
	return print('文本保存完毕')

def to_word(text,root):
	document=Document()
	text_line=text.split('\n')
	cnt=0
	name_list=file_name(root)

	for line in text_line:													#这里的逻辑我自己都看的不忍直视了，求大神优化
		pat_text_pic=re.compile('图片'+str(cnt))							#构建“图片0”的正则式
		if pat_text_pic.search(line):										#发现文章里出现了“图片0”的字样
			document.add_picture(root+name_list[cnt],width=Inches(1.25))	#在word里加入0.jpg
			print("word添加"+name_list[cnt]+"成功")							
			cnt=cnt+1
		else:
			document.add_paragraph(line)
	document.save(root+'content.docx')		


'''-------v1.4版用---------------------
def save_path():
	root=input("输入你的文件要保存的位置。直接按回车，内容默认存至F:\\pic_test")
	if(root==''):
		root='F:\\pic_test\\'
	if(root[-1]!='\\'):
		root=root+'\\'
	return root	
'''
def save_path(title,num):
	title_re=re.sub(r'[\.\!\/_,$%^*(+\"\')]+|[+——()?【】“”！，。？、~@#￥%……&*（）]','',title[0:-5])	#[0:-5]是去除“--知乎”的结尾，正则替换则是去除标点。因为文件名无法用标点符号。
	root='F:\\zhihu_spider\\'+title_re+"\\"+str(num)+'\\'
	return root

def file_name(file_dir):															#考虑到保存下来的图片后缀名不一定一致，需要遍历一遍文件名
	name_list=[]
	for root,dirs,files in os.walk(file_dir[:-1]):									#此时的root是X:\abc\的形式，需要变成X:\abc的形式
		for file in files:
			if not re.search(r'\.txt|\.docx',file):																
				name_list.append(file)												#剔除txt外的子文件名，也就是全部图片的文件名。context.txt的名字会影响lambda算法的排序
	name_list.sort(key = lambda i:int(re.match(r'(\d+)',i).group()))				#字符串数字排序，如果不采用lambda，那么排序可能是['1','11','2']的形式。这个算法好牛逼！

	return name_list		

'''
   print(root) #当前目录路径  
   print(dirs) #当前路径下所有子目录  
   print(files) #当前路径下所有非目录子文件 
'''



def get_url_from_collection(url):
	resp=html_download(url)
	soup=BeautifulSoup(resp,"html.parser")
	data_temp=[]																				#获取收藏里每一个链接
	data_dup=[]																					#知乎里一个href会出现好几次
	data=[]																						#获取收藏里每一个有效链接
	url_origin=soup.find_all(['link','a'])	

	for u in url_origin:
		if 'href' in u.attrs:
			data_temp.append(u.attrs['href'])											
										
	for i in range(len(data_temp)):
		if 'answer' in data_temp[i] and not 'apply' in data_temp[i]:															#收藏有专栏文章和回答两种。专栏文章的url是完整的
			data_dup.append('https://www.zhihu.com'+data_temp[i])
		if 'zhuanlan' in data_temp[i]:
			data_dup.append(data_temp[i])
	data=list(set(data_dup))
	data.sort(key=data_dup.index)			
	print(data)	
	return data
				




def main():
	print("制作：CapGhost\r特别感谢：娜娜&饭团")

	collection_url=input("输入你的知乎收藏文件夹的网址，输入回车则选择缺省网址：\n")
	if(collection_url==''):
		collection_url='https://www.zhihu.com/collection/211012874'
#	root=save_path()								#1.4版用

	url_list=get_url_from_collection(collection_url)
	print("你的收藏里共"+str(len(url_list))+"篇文章，现在开始爬取。。。")
	cnt=0

	for url in url_list:
		print("第"+str(cnt+1)+"篇：")
		data=html_download(url)
		title=get_title(data)
		num=url.split("/")[-1]														#考虑到会收藏同样问题的多个回答，故在路径最后添加一组数字，为网址最后的数字
		pic_url=get_pic_url(data)
		text=get_text(data)
		root=save_path(title,num)														#文件保存在F:\问题（或者文章）名\
		print(root)
		save_pic(pic_url,root)
		save_text(text,root)
		to_word(text,root)
		cnt=cnt+1

	


if __name__ == '__main__':
	main()
	

'''
1.r'\>[^\<]+\<|data-actualsrc=".*?\"'		#注意用问号解决re库默认的贪婪匹配！
2.str.replace('>','').replace('<','\r')	#只是取消">"<"这两个标识符，链式replace即可
3.用正则来处理文本，而不是用BS4通过提取标签的方法来处理
4.str是关键字，用于str()的转换。str也可设为变量，但考虑此函数，不要用这个变量名
5.网上复制代码有风险，特别是python这种游标卡尺的
'''